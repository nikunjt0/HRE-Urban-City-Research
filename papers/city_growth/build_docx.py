"""Build the revised paper as a .docx (research-narrative rewrite).

Fixes vs the old draft (data-audit revision, Aug 2026):
  - privilege analyses restricted to each source's coverage universe (no false zeros);
  - stacked event studies, treatment counts, MDEs/equivalence bounds, exposure timing;
  - no-leakage backcast (coefficients locked on pre-1200 transitions) + spatial CV ceiling;
  - Shapley framed as descriptive/conditional; residual renamed 'unexplained & weakly persistent';
  - Black Death framed as extension of Jedwab-Johnson-Koyama;
  - Buringh imputation quantified; threshold robustness; independence claim qualified.
Output: out/European City Growth Paper (revised).docx
"""
from __future__ import annotations
from pathlib import Path
import pandas as pd, numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
OUT = HERE / "out"
FIG = OUT / "figures"
DOC = OUT / "European City Growth Paper (revised).docx"

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Georgia"; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(7)

FIGN = {"n": 0}
TABN = {"n": 0}


def title(t, sub, author, date):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(22); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sub); r.font.size = Pt(12); r.italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{author}\n{date}").font.size = Pt(11)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a); r.font.size = Pt(15)


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33); r.font.size = Pt(12)


def para(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        p.add_run(bold_lead + " ").bold = True
    p.add_run(text)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        p.add_run(bold_lead + " ").bold = True
    p.add_run(text)


def fig(fname, caption, width=6.1):
    FIGN["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / fname), width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(f"Figure {FIGN['n']}. {caption}")
    r.font.size = Pt(9); r.italic = True


def table(headers, rows, caption, widths=None):
    TABN["n"] += 1
    c = doc.add_paragraph()
    r = c.add_run(f"Table {TABN['n']}. {caption}")
    r.font.size = Pt(9.5); r.bold = True
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htxt); run.bold = True; run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    doc.add_paragraph()


def equation(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.name = "Cambria Math"; r.font.size = Pt(12); r.bold = True


# ================================================================ front matter
title("Cities Grew for Boring Reasons",
      "The Determinants of City Size in the Holy Roman Empire and Central Europe, 1200–1500",
      "Nikunj Tyagi", "August 2026")

h1("Abstract")
para("Why did some medieval cities turn into metropolises while their neighbors stayed villages? "
     "The standard answer credits commercial privileges and institutions: staple rights, chartered "
     "liberties, fairs, market rights. Using two standard reconstructions of European city "
     "populations (2,262 cities, 700–2000) joined to dated privilege records — with each privilege "
     "analysed strictly inside the geographic universe its source actually documents — I test that "
     "answer and find no support for it. Inside the Viabundus network area, towns holding staple "
     "rights were 40% larger than towns without them; Europe-wide, communes were 21% larger than "
     "non-communes. But towns that received a privilege had been rising relative to their peers "
     "for centuries before the grant, and their growth does not accelerate after it: "
     "difference-in-differences estimates around the grant are approximately zero for six "
     "institutional treatments from four sources (staple −7% [−22, +10]; fair −0% [−13, +14]; "
     "charter +0% [−20, +28]; market −3% [−26, +32]; and — extending the test beyond Germany to "
     "the Italian and French commune belt — communal self-government −1% [−9, +7] and "
     "participative institutions −1% [−8, +8]). For the well-powered treatments, effects larger "
     "than +7–14% are excluded; for charters and market rights the intervals remain wide, so "
     "those nulls are imprecise. What the data consistently reject is the traditional reading in "
     "which the grant launches the growth: the ascent begins before the privilege arrives.")
para("What did determine the urban hierarchy of 1500? A Shapley decomposition of predictive R² — "
     "descriptive accounting, not causal attribution — traces 48% of the total variance in 1500 "
     "size to a city's own size in 1200 and 19% to its size in 800 AD, against 2.6% for measured "
     "geography conditional on those sizes. The conditional qualifier matters: geography's small "
     "direct share does not mean geography barely built the hierarchy — water and Roman-era "
     "location did their work early, and the sizes of 800 and 1200 embed it. On the subsample "
     "where all four privilege statuses are observable, privileges add about four points of R² to "
     "a geography-plus-size model. The remaining ~31% is unexplained and only weakly persistent: "
     "it barely shrinks when institutions are added, and a city's growth residual in one century "
     "correlates at −0.10 to −0.16 with the next — consistent with transient shocks plus "
     "rounding-induced measurement bounce, though not proof of pure randomness. The one lever "
     "geography still pulled after 1200 was navigable water, worth about 5–6% extra growth per "
     "century (+15% by 1500, CI [+7%, +23%]), and central to the reshuffle after the Black Death.")
para("These facts compress into a three-parameter law of motion. Handed only the populations of "
     "1200 and a water-access dummy, the law predicts the city sizes of 1500 with R² = 0.56 out of "
     "sample, places 81% of cities within a factor of two, and identifies eight of the ten largest "
     "cities of 1500. Crucially, this survives a strict no-leakage test: coefficients estimated "
     "solely on the 800→1200 transitions — locked before the forecast window opens — still "
     "deliver R² = 0.50 (0.56 after a single overall level correction), the same rank correlation, "
     "and the same eight of the ten largest cities. Flexible machine-learning models given every "
     "feature knowable in 1200 reach only ≈0.60 under spatially blocked cross-validation, so the "
     "law sits near the information ceiling. Meanwhile the same law predicts essentially none of "
     "which individual cities would rise or fall (growth R² = 0.02). Both numbers are the point: "
     "the position of the urban hierarchy was highly predictable; movement within it was not.")

# ================================================================ 1
h1("1. The Question, and the Order the Answers Came In")
para("Medieval history offers a ready explanation for why Hamburg, Nürnberg, or Danzig became great: "
     "they won privileges. A staple right forced passing merchants to unload and offer their goods "
     "for sale. A charter freed townsmen from feudal obligations. A fair concentrated the region's "
     "trade inside the walls twice a year. These grants were fought over, paid for, and celebrated, "
     "so it is natural to assume they mattered. The assumption is testable, and this paper tests it.")
para("The paper is organized in the order the research actually unfolded, because each result forced "
     "the next question:")
bullet("Is growth predictable from size? If big cities systematically out-grew small ones "
       "(or vice versa), that mechanism would come first. It turns out neither is true — growth is "
       "almost unrelated to starting size — which means the hierarchy persists by default and the "
       "interesting question becomes what set it up.", "Step 1 (§3).")
bullet("Did privileges cause growth? The raw correlations say yes; a comparison that respects "
       "timing says no.", "Step 2 (§4).")
bullet("Build an explicit benchmark for what each city ‘should’ have weighed in 1500 given its 1200 "
       "size, so that over- and under-performance can be measured city by city rather than "
       "asserted.", "Step 3 (§5).")
bullet("Decompose the 1500 hierarchy into candidate sources — inherited size, deep origins, "
       "geography, and (back in for a fair hearing) institutions — and measure what is left "
       "over.", "Step 4 (§6).")
bullet("Identify the one geographic factor that still moved cities: navigable water.", "Step 5 (§7).")
bullet("Compress everything into a single equation, and — the acid test — ask how well that "
       "equation actually predicts the map of 1500 from the map of 1200.", "Step 6 (§8).")
para("Section 9 repeats the core results on the alternative population reconstruction and "
     "stress-tests the data's imputation and rounding; §10 states explicitly what is new here; "
     "§11 concludes.")

# ================================================================ 2
h1("2. Data")
para("Medieval city populations are not measured; they are reconstructed from fragments — tax "
     "rolls, hearth counts, militia lists, church registers — converted into inhabitants with "
     "demographic multipliers. Because this is an estimation exercise, I run everything on the two "
     "most widely used reconstructions. They are built by different scholars but from overlapping "
     "source bases — Buringh explicitly revises and extends Bairoch, and about 7% of the medieval "
     "observations used here cite Bairoch as their direct source — so they are best read as "
     "partially overlapping reconstructions, not fully independent measurements:")
bullet("Bairoch, Batou & Chèvre (1988), La population des villes européennes de 800 à 1850: the "
       "foundational compilation — roughly 2,200 cities at hundred-year intervals.")
bullet("Buringh (2021), The Population of European Cities from 700 to 2000: a later, georeferenced "
       "reconstruction of 2,262 cities that revises Bairoch with newer scholarship and attaches "
       "coordinates, elevation, and a transport classification (river, coastal, landlocked).")
para("Where both sources cover a city in 1500, their log populations correlate at r = 0.96 — close "
     "agreement, but not identical, which is what makes the robustness check in §9 meaningful. "
     "Buringh is the primary panel throughout (it carries the geography); Bairoch is the check. The "
     "sample is the Holy Roman Empire and its neighbours — Germany, Italy, France, the Low "
     "Countries, Poland, Hungary, Switzerland, Austria — restricted to cities of at least 1,000 "
     "inhabitants: 1,259 cities in all.")
para("Two data caveats matter throughout. First, both reconstructions round populations to "
     "convenient steps (1,000, 2,000, 4,000…). Rounding injects spurious jumps into measured "
     "growth — a city drifting from 1,400 to 1,600 people appears to leap from 1,000 to 2,000. "
     "Part of what no model can explain in §6 is this measurement noise, and I will show a "
     "diagnostic that says so. Second, Buringh flags the nature of every city-year value, and in "
     "the 1200–1500 analysis sample 81% of observations are imputed or proxied (95% in 1200): "
     "city- and time-specific estimates, not direct documentary counts. Too few purely documentary "
     "observations survive in this era (≈40 cities) to re-run the analysis on them alone; the "
     "checks that are feasible — swapping in Bairoch's numbers (§9) and raising the population "
     "threshold to 5,000 and 10,000 (§9.2) — leave every core result qualitatively unchanged.")
para("Privileges and institutions come from two further sources:")
bullet("Viabundus (Holterman et al.), a georeferenced reconstruction of the late-medieval road, "
       "river and sea network of Northern Europe, with dated staple rights, fairs, and tolls "
       "(156 staples, 773 fairs).")
bullet("Cantoni, Mohr & Weigand (2020), dated town charters (Stadtrecht) and market rights "
       "(Marktrecht) for the 2,390 towns of the Deutsches Städtebuch, with legal families "
       "(Magdeburg law, Lübeck law, …).")
para("Three further sources extend the institutional record beyond the German privilege "
     "datasets — each again used strictly inside its own universe:")
bullet("Bosker, Buringh & van Zanden (2013), a 792-city European/Mediterranean panel whose "
       "commune variable records communal self-government by century, 800–1800. This is the "
       "Europe-wide analogue of a town charter, and it covers exactly the region the German "
       "sources cannot: Italy, France, Austria, Switzerland, Hungary, the Low Countries.")
bullet("Wahl (2015), participative political institutions (council elections, guild "
       "participation, burgher representation) for 325 central-European cities including "
       "Austria and Switzerland, by century 800–1800.")
bullet("Krauer & Schmid (2022), a geocoded digitization of Biraben's plague-outbreak "
       "inventory: 11,180 dated outbreak records 1346–1900, of which 338 place-level records "
       "fall in the first Black Death wave (1347–1352). Used in §7 to test the plague "
       "mechanism directly.")
para("The Bosker and Wahl institutions are century-resolution status panels rather than dated "
     "grants: if an institution is first observed at a census year, adoption occurred sometime "
     "in the preceding century, and it is dated to that century's midpoint — equivalent to a "
     "mid-century grant with ~50 years of exposure before the first treated census, well within "
     "the exposure range of the exactly-dated privileges.")
para("A blank in a privilege source can mean two very different things: 'this town did not have "
     "the privilege' or 'this town is outside the area the source documents.' Conflating the two "
     "manufactures false zeros. Viabundus maps the road, river and sea network of northern and "
     "central Europe — its nodes stop near latitude 49°N, so Italy, Austria, Switzerland, Hungary "
     "and most of France lie beyond the map's edge; the Deutsches Städtebuch covers Germany in its "
     "1937 borders, so Vienna, for example, simply is not in it. Every privilege analysis in this "
     "paper is therefore restricted to the universe in which absence is actually observable: "
     "staple and fair tests to cities within 25 km of a Viabundus network node (432 of the 1,259 "
     "sample cities — the median covered city sits about 1 km from a node, the median excluded "
     "city hundreds of km away, so the cutoff is not sensitive), and charter and market tests to "
     "cities matched to a Städtebuch town (name match within 60 km, or coordinates within 5 km); "
     "likewise, commune and participative-institution tests are restricted to cities matched to a "
     "Bosker or Wahl city (within 8 km), and plague exposure to cities near a geocoded outbreak "
     "record. A city outside a source's universe is treated as missing data for that variable — "
     "never as an untreated zero.", "Coverage discipline.")

# ================================================================ 3
h1("3. Step 1: Growth Was Not Predictable from Size")
para("A city's size in 1500 is the compound of everything that happened to it before. If some "
     "cities ended up vastly larger than others, then either they grew faster during 1200–1500, or "
     "they started ahead and coasted. So the first thing to establish is the basic character of "
     "growth itself.")
para("Figure 1 plots each city's growth over 1200–1500 against its size in 1200. If size bred "
     "growth (big cities compounding their advantage), the cloud would slope up. If small towns "
     "systematically caught up, it would slope down.")
fig("fig_gibrat.png",
    "City growth 1200→1500 against starting size. Starting size explains ~2% of growth.")
para("The cloud is nearly flat: starting size explains about 2% of subsequent growth. This is "
     "Gibrat's law — growth roughly independent of size — and it has a consequence that shapes "
     "everything after. If growth does not depend on size, then a city 10× larger than its "
     "neighbor in 1200 tends to stay roughly 10× larger, because both draw growth from the same "
     "distribution. Rankings persist by default, without any mechanism favoring the big.")
para("Individual cities still followed wildly different paths — the vertical spread in Figure 1 is "
     "huge. But those differences were not systematically related to where a city started. That "
     "re-aimed the whole project: instead of hunting for a recipe that made cities grow, I needed "
     "to ask (a) what created the initial differences that later growth carried forward, and "
     "(b) whether policies and privileges let particular cities beat the average fate their size "
     "implied. Those are Steps 4 and 2 respectively.")

# ================================================================ 4
h1("4. Step 2: Privileges Fail the Causal Test")
para("Start with the most prized privilege, the staple right (Stapelrecht), which forced merchants "
     "passing through to unload and offer their cargo for sale. Inside the Viabundus universe, "
     "towns holding a staple right were on average 40% larger than towns without one (two-way "
     "fixed effects); towns with fairs were 26% larger. (Ignoring coverage inflates these gaps to "
     "53% and 41%, because every uncovered southern city gets miscounted as unprivileged — one "
     "measure of how much the false-zero problem distorts.) Historians have long read gaps like "
     "these as evidence that privileges built cities.")
para("The problem is timing. Staple rights were granted to towns that had already become important "
     "junctions; fairs were awarded to towns that already drew traders. A cross-sectional gap "
     "cannot distinguish ‘the privilege made the town big’ from ‘the town was big, so it got the "
     "privilege.’ Separating those requires following towns through time around the moment of the "
     "grant. Concretely, for every town that received a privilege in some year, I:")
bullet("recorded its population in the centuries before and after the grant year;")
bullet("matched it to comparison towns of similar size, in the same sample, with the same water "
       "access, that did not receive the privilege in that window;")
bullet("computed the difference-in-differences: (growth of the granted town after − before) minus "
       "(growth of its comparison towns after − before). If the privilege caused growth, this "
       "quantity is positive.")
para("Three things make the design credible — and honest about its limits. First, coverage: each "
     "privilege is tested only inside its source's documented universe (§2), with treatment counts "
     "stated: 88 towns with dated staples, 143 with dated fairs, 270 with dated charters, 208 with "
     "dated market grants (fewer enter each regression once three consecutive centuries of "
     "population data are required). Second, a stacked event study around the grant centuries — "
     "not-yet-treated and never-treated towns as controls, cohort-stacked fixed effects, "
     "city-clustered errors — shows the full trajectory rather than a single difference (Figure "
     "3). Third, the same machinery applied to four privileges from two independently assembled "
     "sources tells one consistent story. Figure 2 shows the headline comparison.")
fig("fig_causal_consolidated.png",
    "Raw size gaps (red) versus the matched difference-in-differences estimate (blue) for six "
    "institutional treatments from four sources, with 95% bootstrap confidence intervals and "
    "treatment counts. No estimate rejects zero; the staple, fair, commune, and "
    "participative-institution intervals are tight, the charter and market intervals are wide.")
para("The raw gaps do not survive the timing-aware comparison: staple −7% [−22, +10], fair −0% "
     "[−13, +14], town charter +0% [−20, +28], market right −3% [−26, +32]. These are null "
     "results of two different strengths, and the difference matters. For the two well-powered "
     "Viabundus privileges, the intervals exclude effects larger than +10% (staple) and +14% "
     "(fair) — small against raw gaps of +40% and +26%, and against the minimum detectable "
     "effects at 80% power (+28% and +21%). A confidence interval that crosses zero means 'cannot "
     "reject zero', not 'proved zero'; for staples and fairs, though, anything close to the "
     "association historians cite is affirmatively ruled out. For charters and market rights the "
     "intervals are wide ([−20, +28] and [−26, +32]) and the minimum detectable effects large "
     "(+41% and +51%): there the honest statement is that no effect is detected but a modest "
     "positive one cannot be excluded.")
fig("fig_event_study.png",
    "Stacked event studies around the grant century (event time −1 = last pre-grant census, "
    "reference), for all six institutional treatments. Towns that will receive a privilege are "
    "already rising relative to controls two centuries before the grant, and the ascent "
    "continues at the same slope after it — no break appears at the grant date, in Germany or "
    "beyond it.")
para("Figure 3 shows why the story is selection rather than treatment. Towns that would receive a "
     "staple right sat roughly 19 log points below their eventual relative level two centuries "
     "before the grant and climbed steadily through it; the climb after the grant continues the "
     "pre-grant slope, with no kink at the grant date. The privilege marks the middle of a long "
     "ascent, not the start of one. This also disciplines how the pre-trends should be described: "
     "they are not flat in levels — treated towns were rising, exactly as 'privileges were "
     "awarded to rising towns' predicts — but the growth rate does not change when the grant "
     "arrives, which is what the difference-in-differences nets out.")
para("Do these nulls generalize beyond Germany — or are they an artifact of testing only where "
     "the German sources reach? Two Europe-wide institutions answer this. Communal "
     "self-government (Bosker et al.), the closest continental analogue of a town charter, is "
     "observed for 391 matched cities including Italy, France, Austria, Switzerland and "
     "Hungary, with 243 dated adoptions. The pattern is identical and more precise than any "
     "German privilege: communes were 21% larger than non-communes in the naive comparison, "
     "but the matched difference-in-differences is −1% [−9, +7] — the tightest interval of "
     "the six treatments; self-government effects larger than +7% are excluded, against a "
     "minimum detectable effect of +12%. Participative institutions (Wahl: council "
     "elections, guild participation, burgher representation; 207 dated adoptions including "
     "Austrian and Swiss towns) give −1% [−8, +8]. The event studies (Figure 3, right panels) "
     "show the same signature as the German privileges: future communes were rising for two "
     "centuries before self-government arrives, and do not accelerate afterward. The Italian "
     "and French communes so celebrated in the historiography were, like the German charters, "
     "milestones on ascents already under way.")
para("A century-grid panel has a mechanical exposure problem: a staple granted in 1202 enjoys 98 "
     "years of exposure before the next census, one granted in 1297 only three. If privileges "
     "worked slowly, late-century grants would bias the estimates toward zero. They do not drive "
     "the result: restricting to grants in the first half of their century (at least 50 years of "
     "exposure) moves the estimates to staple −5%, fair +1%, charter −15%, market −6% — the same "
     "picture. (The commune and participative-institution treatments are midpoint-dated by "
     "construction, ~50 years of exposure each, squarely in this range.)")
para("Growth in the century after a grant is statistically indistinguishable from growth in the "
     "century before it, across six institutional treatments, four sources, and both sides of "
     "the Alps. For the well-powered treatments — staples, fairs, communes, participative "
     "institutions — effects above +7–14% are excluded outright. This parallels Bosker, Buringh "
     "& van Zanden (2013), whose ‘bishopric advantage’ evaporates once city fixed effects absorb "
     "the fact that bishops sat in already-important places.",
     "Finding 1: privileges were badges of arrival, not engines.")
para("The Städtebuch record makes the logic vivid. Cologne, the empire's largest city (40,000 in "
     "1200), was a Roman colonia with no formal charter bestowal in the Städtebuch record — it "
     "held city status long before formal charters became the instrument. Hamburg (2,000 in 1200) "
     "received charter and staple and grew fifteen-fold. Mainz held both a charter (1300) and a "
     "staple right and shrank from 9,000 to 6,000. Knowing a town's privileges tells you little "
     "about its fate; §5 makes that precise.")

# ================================================================ 5
h1("5. Step 3: What a City ‘Should’ Have Weighed — the Implied-Size Benchmark")
para("Claims like ‘Hamburg over-performed’ or ‘Mainz under-performed’ are empty until "
     "‘performance’ is defined against a benchmark. Here is the one used for the rest of the "
     "paper. Take the 1,092 cities observed in both 1200 and 1500 and fit the average relationship "
     "between the two dates by least squares:")
equation("log P₁₅₀₀ = 1.34 + 0.86 · log P₁₂₀₀ + 0.14 · water")
para("This line is what ‘the size its 1200 population implies’ means: it is the 1500 size of the "
     "average city that started at a given 1200 size (with a modest bonus if it sat on navigable "
     "water — the water term is developed in §7). A slope of 0.86, slightly below 1, says the "
     "hierarchy compressed a little: a city 10× larger than another in 1200 was typically about "
     "7× larger in 1500. A city's performance is its residual — how far its actual 1500 "
     "population sits above or below the line, in log points. Figure 4 draws this.")
fig("fig_implied_size.png",
    "The implied-size benchmark. Each dot is a city; the blue line is the fitted average "
    "1200→1500 relationship for cities on water (green dashed: landlocked). A city's "
    "performance is its vertical distance from the line (red segments; values in log points).")
para("Now individual cities can be read quantitatively. Hamburg's residual of +2.27 log points "
     "means it ended at e²·²⁷ ≈ 9.7 times its implied size; Amsterdam (+1.96) at 7×; Antwerp "
     "(+1.93) at 7×; Danzig (+1.48) at 4×. Mainz (−0.63) ended at barely half the size its 1200 "
     "population implied. Cologne (+0.10) and Vienna (+0.02) grew almost exactly as their starting "
     "sizes dictated — impressive in absolute terms, unremarkable relative to the benchmark.")
para("Ranking all cities by this residual (Figure 5) reveals the geography of over- and "
     "under-performance: the over-performers are almost entirely Atlantic and North Sea ports; the "
     "under-performers are Mediterranean and landlocked interior towns. That pattern is Step 5's "
     "subject.")
fig("fig_performers.png",
    "Cities ranked by performance against the implied-size benchmark. Over-performers (top) are "
    "dominated by Atlantic and North Sea ports; under-performers (bottom) by Mediterranean and "
    "landlocked towns.")

famous = pd.read_csv(OUT / "famous_cities.csv")
rows = []
for _, r in famous.iterrows():
    implied = int(round(r["pred1500"], -2))
    ratio = np.exp(r["resid"])
    if not r["in_stadtebuch"]:
        charter = "n/c"
    elif pd.isna(r["charter_year"]):
        charter = "—"
    else:
        charter = str(int(r["charter_year"]))
    if pd.isna(r["staple"]):
        staple = "n/c"
    else:
        staple = "yes" if r["staple"] == 1 else "—"
    rows.append([r["city"], f"{int(r['pop1200']):,}", f"{implied:,}",
                 f"{int(r['pop1500']):,}", f"{ratio:.1f}×", charter, staple])
table(["City", "Pop 1200", "Implied 1500", "Actual 1500", "Actual / implied", "Charter", "Staple"],
      rows,
      "Twenty-one major cities against the implied-size benchmark. ‘Implied 1500’ is the fitted "
      "value from the benchmark regression; the ratio is the city's over- or under-performance. "
      "Privilege columns distinguish an observed absence (‘—’: the city is inside the source's "
      "coverage area and no grant is recorded) from missing coverage (‘n/c’: the city lies "
      "outside the source's universe — the Städtebuch covers 1937 Germany, Viabundus the "
      "northern trade network — so its status is unobserved, not absent).")
para("The table rewards a slow read. The explosive over-performers (Hamburg 9.7×, Amsterdam 7.1×, "
     "Antwerp 6.9×, Utrecht 5.7×, Ulm 5.2×) are ports and river towns of the Atlantic–North Sea "
     "world. The Roman-era giants tracked their benchmarks — and their charter entries need care. "
     "Cologne and Augsburg are inside the Städtebuch universe and record no formal charter "
     "bestowal: they held city status before formal charters became the instrument (Augsburg's "
     "civic constitution was already being confirmed under Barbarossa in 1152, and its famous "
     "Stadtbuch of 1276 codified existing law — the Städtebuch classes both under ‘city character "
     "prior to town charter’). Vienna and Basel, by contrast, are simply outside the source's "
     "coverage (‘n/c’) — their blank is missing data, not an absent charter. And privileges "
     "scatter across the whole range of outcomes: Mainz held charter and staple and finished at "
     "0.5×; Nürnberg held a charter and finished at 4.2×; Augsburg, with no formal charter, "
     "finished at 1.9×. The residual — not the privilege list — is the quantity with structure "
     "in it.")

# ================================================================ 6
h1("6. Step 4: Where the 1500 Hierarchy Came From")
para("If privileges did not build the hierarchy, what did? This section decomposes the variation in "
     "log city size in 1500 across candidate explanations. The candidates are not arbitrary; they "
     "are the four theories on the table at this point in the argument:")
bullet("the path-dependence channel: a city's own size in 1200, at the start of the "
       "study window. Step 1 showed growth is unrelated to size, which makes early size a "
       "candidate to dominate by pure carry-forward.", "Inherited size —")
bullet("its size already in 800 AD, the Carolingian era — the earliest layer the data "
       "supports. This separates medieval momentum from origins laid down before the Middle Ages "
       "(Roman roads, bishoprics, river fords).", "Deep origin —")
bullet("river and coastal access, sea basin, elevation — the ‘where it sits’ "
       "channel, which operates whether or not anyone grants anything.", "First-nature geography —")
bullet("staple, fair, charter, and market rights held by 1500. Step 2 showed no growth "
       "break at their grants, but they stay in the race here as predictive markers — measured "
       "only on the subsample where their status is actually observable (§2), so that no false "
       "zeros dilute them and the unexplained remainder cannot be blamed on their omission.",
       "Institutions —")
para("The decomposition proceeds in two transparent steps (Figure 6).")
h2("6.1 Step one: the R² ladder")
para("First, ask how much of the variation in 1500 size each candidate explains on its own, by "
     "regressing log 1500 size on each block separately. Geography alone: 9%. Size in 800 alone: "
     "40% (n = 551). Size in 1200 alone: 56% (n = 1,092). Geography + 800 + 1200 jointly: 69%. "
     "Institutions alone, on their observable universe: 24% (n = 222) — note that an uncorrected "
     "analysis, with every uncovered city miscounted as unprivileged, would put this at ~5%; "
     "honest coverage makes institutions look more predictive, not less, because the comparison "
     "pool is no longer polluted. Adding institutions to geography + 1200 size on that same "
     "universe lifts R² from 0.57 to 0.61.")
fig("fig_decomp_build.png",
    "Left: variance in 1500 city size explained by each candidate alone, then jointly (R²). "
    "Right: the Shapley split of the joint R² on the full sample (n = 551); the ~31% no model "
    "explains is shown in grey. Shares are descriptive splits of predictive R², not causal "
    "attributions.")
h2("6.2 Step two: splitting shared credit fairly (Shapley)")
para("The solo numbers overlap: 9 + 40 + 56 sums to far more than the joint 69, because the "
     "candidates are correlated — cities big in 1200 tended to be big in 800, and both tend to sit "
     "on water. Whichever variable enters a regression first soaks up the shared credit. The "
     "Shapley decomposition (Shorrocks 1982; Israeli 2007) removes that arbitrariness: it computes "
     "each candidate's marginal contribution to R² under every possible order of entry and "
     "averages them. Nothing about the method favors any candidate — the same arithmetic that "
     "credits inherited size would have credited institutions, had the data pointed that way.")
para("The result on the full sample (n = 551), as shares of the total variance in 1500 size: "
     "inherited size (1200) 48%; deep origin (800) 19%; geography 2.6%; unexplained 31%. "
     "Institutions cannot enter this split honestly (their status is unobserved outside their "
     "sources' coverage), so they are bounded on the privilege-observable universe (n = 196): "
     "there, a three-way split gives inherited size 44 points of variance, institutions 12, "
     "geography 6, with 39 unexplained. The same exercise on the Europe-wide Bosker universe "
     "(n = 367, including the Italian and French commune belt) is strikingly parallel: "
     "inherited size 42 points, communal self-government 11, geography 4. Even taking these "
     "11–12-point shares at face value, institutions are markers, not movers: Step 2 found no "
     "growth break at any of the six grants, and §8 shows they add little to actual "
     "forecasting accuracy. North and south of the Alps alike, the institution flags the kind "
     "of city that grows; it does not make cities grow.")
para("Two glosses keep the headline numbers honest. First, these are shares of predictive R² — a "
     "descriptive accounting of where the variance sits, not a causal attribution; nothing here "
     "says 46% of a city's 1500 population was 'caused by' its 1200 size. Second, geography's "
     "2.6% is a conditional share: conditional on the size a city had already reached by 1200, "
     "measured geography adds only ~2.6 points of explanatory power. That is emphatically not "
     "the claim that geography barely mattered in creating the hierarchy. The 1200 sizes are "
     "themselves the cumulative product of centuries of geography acting through trade — Bosker, "
     "Buringh & van Zanden's evidence that transport geography seeds cities is fully compatible "
     "with this result — so the small conditional share means geography's work was largely done "
     "and banked by 1200, its effects carried forward inside the inherited sizes. The "
     "decomposition separates 'still operating after 1200' from 'already embedded'; it cannot, "
     "and does not, measure geography's total historical contribution.")
h2("6.3 Interrogating the unexplained 31%: weakly persistent, not proven luck")
para("A skeptical reader should ask two questions about the ~31%: is it factors I did not choose "
     "to include, and is it genuinely random? The evidence answers the first firmly and the "
     "second only partly — which is why this paper calls it ‘unexplained and weakly persistent "
     "variation’ rather than ‘luck.’")
bullet("It is not the institutions. On the universe where all four privilege statuses are "
       "observed, adding all four blocks to the geography-plus-size model moves the unexplained "
       "share from 43.3% to 38.9% — institutions were hiding about 4 points, not 30.", "(i)")
bullet("It is not any stable hidden factor. This is the sharp test. Any durable city trait "
       "omitted from the model — an entrepreneurial elite, a lucrative relic, an unusually able "
       "council — would push the same city above its benchmark century after century, making "
       "growth residuals persist. They do not: the correlation between a city's growth residual "
       "in one century and the next is −0.10 (1200s→1300s) and −0.16 (1300s→1400s).", "(ii)")
bullet("But weak persistence does not prove randomness. Measurement error alone induces mean "
       "reversion (populations rounded to 1,000/2,000/4,000 steps create exactly the slight "
       "negative autocorrelation in (ii)), and genuinely systematic causes can be transient "
       "rather than permanent: a war, a fire, a banking boom, a rerouted trade artery need not "
       "repeat next century. The residual is therefore an upper bound on true historical luck, "
       "and a mixture in unknown proportions of transient real shocks — plagues, sieges, "
       "dynastic accidents, industry booms — and recording noise. What it is not, on the "
       "evidence of (i) and (ii), is a stable factor waiting to be measured.", "(iii)")
para("How deep does the path dependence run? Figure 7 traces the lock-in backward: a city's size "
     "in 800 AD — the age of Charlemagne — still explains 40% of its size in 1500, seven hundred "
     "years later. The urban map of the Holy Roman Empire in 1500 was, to a first approximation, "
     "drawn before the Middle Ages began.")
fig("fig_persistence_depth.png",
    "Depth of persistence: variance in 1500 size explained by size at each earlier date. Even "
    "800 AD explains 40%.")
para("About two-thirds of a city's 1500 size was already determined by its own earlier size; no "
     "measured factor, and no stable unmeasured factor, accounts for more than a sliver of the "
     "rest; and the remainder is weakly persistent — transient shocks and measurement noise, "
     "with true luck as a large but not separately identified component.", "Finding 2:")

# ================================================================ 7
h1("7. Step 5: Water — the One Lever Geography Still Pulled")
para("One geographic factor did keep moving cities after 1200: navigable water. Regressing "
     "1200→1500 growth on water access (controlling for starting size), cities on a river or "
     "coast grew +0.143 log points more — about 15% larger by 1500, bootstrap 95% CI [+7%, +23%] "
     "— than comparable landlocked cities, roughly +5–6% per century compounding. Unlike "
     "privileges, this factor cannot be endogenous to success: a town cannot petition its way to "
     "the coast. Where the privilege gaps collapsed under the causal test, the water premium has "
     "nothing to collapse into — the geography was there first.")
para("The premium was not constant — it was switched on late. Figure 8 resolves the coastal "
     "premium by sea basin and century. In the plague century (1300–1400), coastal cities did "
     "worse, Mediterranean ports especially; after 1400, every basin turns strongly positive as "
     "long-distance trade reorganized around the Atlantic, North Sea, and Baltic. What the data "
     "show in reduced form: the cities that emerged from the plague-and-recovery window "
     "permanently larger were disproportionately on water (64% of winners vs 57% of losers; net "
     "1300→1500 gain on water +0.084 log, p = 0.008, controlling for pre-plague size).")
fig("fig_water_timing.png",
    "The water premium by sea basin and century. Negative for Mediterranean ports in the plague "
    "century; strongly positive everywhere after 1400.")
para("This pattern is best read as an extension of Jedwab, Johnson & Koyama, not a rival "
     "novelty: they establish that European cities recovered from the Black Death toward "
     "places with stronger fixed factors of production. The refinement offered here is that, "
     "within this Central European sample, the operative fixed factor was navigable water — "
     "the premium is concentrated in the recovery century and shows up in every basin — which "
     "ties their recovery mechanism to the specific geography that also dominates the "
     "over-performer list of §5.")
h2("7.1 Testing the mechanism with plague-arrival data")
para("Maritime plague entry is the natural reading of the 1300s dip — Messina and Marseille "
     "were famously the disease's ports of entry — and rather than leave it as an assertion, "
     "it can be confronted with data: the Krauer & Schmid digitization of Biraben's outbreak "
     "inventory supplies 338 geocoded, place-level outbreak records from the first wave "
     "(1347–1352), matched here to sample cities within 10 km. The test comes with a built-in "
     "health warning. Biraben's inventory is a chronicle compilation, not a census of "
     "outbreaks: absence of a record is weak evidence of absence, recording probability rises "
     "with a town's documentation density (the match-rate table in the appendix shows the "
     "Low Countries and Hungary at implausible zeros, a known undercoverage), and initial size "
     "is controlled throughout for exactly that reason.")
para("The results support the arrival leg of the mechanism and honestly fail to identify the "
     "rest (Figure 9). Arrival: coastal cities were 7 percentage points and river cities 6 "
     "points more likely to record a first-wave outbreak than landlocked cities of the same "
     "size (p = .02 and .001; 142 of 1,174 sample cities hit) — the wave demonstrably "
     "travelled the water network. Recorded arrival years, however, are indistinguishable "
     "across coastal, river, and landlocked cities (all ≈1348): at annual resolution over a "
     "five-year, continent-crossing wave, chronicle dates cannot rank who was struck first. "
     "Impact and recovery: a recorded hit does not predict slower 1300→1400 growth — the "
     "coefficient is actually positive, which is what selection on documentation produces "
     "(thriving towns keep better records) when a century grid averages the 1347–52 crash "
     "with fifty years of rebound; and the recovery-century hit×water interaction is "
     "positive but insignificant (+0.07, p = .47). The honest summary: the exposure geography "
     "of the Black Death confirms that water carried the shock, while the damage-and-rebound "
     "channel is not identifiable from chronicle records at century resolution — it would "
     "need annual population or mortality series.")
fig("fig_plague_mechanism.png",
    "The plague mechanism against arrival data (Biraben, first wave 1347–52). Left: water "
    "cities are significantly more likely to record an outbreak (size-controlled); arrival "
    "years do not discriminate. Right: century-resolution impact and recovery coefficients — "
    "the positive 'hit' coefficient reflects selection on documentation, not plague benefiting "
    "cities.")
para("This is also the pattern behind Figure 5 and Table 1: the over-performers were Atlantic and "
     "North Sea ports because water is where the residual structure lives — the one systematic "
     "escape from the implied-size benchmark.")

# ================================================================ 8
h1("8. Step 6: One Equation — and How Well It Actually Predicts")
para("Steps 1–5 leave three moving parts: a city's own past size, water, and century-wide shocks. "
     "The natural summary is the growth regression itself, written as a law of motion. Pooling all "
     "1200→1300, 1300→1400 and 1400→1500 transitions:")
equation("log P(i, t+1) = a(t) + 0.915 · log P(i, t) + 0.060 · water(i) + ε(i, t)")
table(["term", "estimate", "meaning"],
      [["a(t): century shocks", "+0.88 (1200s), +0.35 (1300s), +1.03 (1400s)",
        "tides common to all cities: expansion, plague century, recovery boom"],
       ["persistence (0.915)", "0.915 per century",
        "a city next century ≈ its current size; deviations decay with a half-life of ~800 years"],
       ["water (0.060)", "+6% per century", "the steady geographic nudge from §7"],
       ["ε: noise, σ = 0.40", "σ = 0.403 log/century",
        "idiosyncratic shocks — 5× the size of the systematic pull (κ/σ = 0.21)"]],
      "The estimated law of motion, term by term.")
para("Nothing here is exotic — it is the regression from Step 1 with the water term from Step 5 "
     "and century intercepts, rearranged. The draft question any reader should ask is: does this "
     "equation actually predict city sizes, or is it a decoration? The test: hand the equation "
     "only the populations of 1200 and each city's water dummy, iterate it forward three "
     "centuries (no noise, no peeking), and compare the predicted 1500 against the actual 1500 "
     "for all 1,092 cities. Table 3 reports the accuracy alongside the benchmarks it has to beat.")
table(["predictor of 1500 sizes (using only 1200 information)", "R² (log size)", "rank corr.",
       "median miss", "within ×2", "top-10 found"],
      [["guess the sample average for every city", "0.00", "—", "×1.63", "62%", "0/10"],
       ["frozen map: predict 1500 size = 1200 size", "0.39", "0.69", "×1.50", "60%", "8/10"],
       ["the equation, without the water term", "0.55", "0.69", "×1.49", "80%", "8/10"],
       ["the equation", "0.56", "0.69", "×1.56", "81%", "8/10"],
       ["the equation, privilege-universe cities only (n = 196)", "0.51", "0.66", "×1.60", "80%", "5/10"],
       ["the equation + all four privileges (same 196 cities)", "0.55", "0.69", "×1.63", "78%", "6/10"]],
      "Prediction accuracy for 1500 city sizes. ‘Median miss’: median multiplicative error "
      "(×1.56 = typical prediction off by 56%). ‘Top-10 found’: how many of the ten largest "
      "actual cities of 1500 appear in the predicted top ten. Privileges are evaluated only on "
      "the cities where all four statuses are observable, against the equation on the same "
      "cities.")
para("Three readings. First, the equation predicts the hierarchy well: R² = 0.56 on log sizes, "
     "81% of cities within a factor of two, and eight of the ten largest cities of 1500 correctly "
     "identified from 1200 data alone (it misses Florence — whose banking-driven surge is exactly "
     "the kind of idiosyncratic rise the law calls noise — and Bordeaux). Second, this is not "
     "city-level overfitting: refitting the law on a random half of the cities and predicting the "
     "other half, 200 times over, gives out-of-sample R² = 0.557 [0.508, 0.605] — identical to "
     "in-sample. Third, privileges carry modest predictive content as markers — about four "
     "points of R² on the cities where they are measured (0.51 → 0.55), consistent with their "
     "Shapley share in §6 — but they do not materially improve identification of the top tier, "
     "and Step 2 showed this content is selection, not treatment: the badge predicts because it "
     "was pinned on winners.")
fig("fig_prediction.png",
    "Left: 1500 sizes predicted by the equation from 1200 data versus actual 1500 sizes "
    "(R² = 0.56; dashed line = perfect prediction). Right: the same equation's predicted versus "
    "actual growth (R² = 0.02).")
para("The right panel of Figure 10 is the other half of the finding. Asked to predict growth — "
     "which cities would rise or fall relative to their start — the same equation manages "
     "R² = 0.02. There is no contradiction: position is predictable because it is inherited; "
     "movement is unpredictable because the noise term (σ = 0.40) dwarfs the systematic pull. "
     "Both facts are the finding. Any book promising the recipe that made particular cities surge "
     "is fitting stories to what the data says is mostly a random draw; meanwhile the boring "
     "regularity — tomorrow's map ≈ today's map plus a water tilt — predicts remarkably well.")
h2("8.1 A stricter test: coefficients locked before 1200")
para("The headline exercise above still contains a subtle temporal leak, and a referee should "
     "catch it: the features are 1200-vintage, but the coefficients were estimated on the "
     "1200→1500 transitions themselves — the fit has seen the outcomes it predicts, and "
     "splitting cities into halves protects only against city-level overfitting, not against "
     "that. Buringh reaches far enough back to close the leak properly: estimate the law solely "
     "on the 800→900, 900→1000, 1000→1100 and 1100→1200 transitions (2,965 city-century "
     "transitions), lock every coefficient, and only then iterate forward from 1200. The "
     "pre-1200 world yields nearly the same law — persistence 0.973, water +0.045 per century — "
     "and Table 4 shows what it forecasts.")
table(["predictor (trained only on 700–1200 data)", "R² (log size)", "rank corr.",
       "median miss", "within ×2", "top-10 found"],
      [["pre-1200 law, strict ex-ante", "0.50", "0.69", "×1.68", "76%", "8/10"],
       ["pre-1200 law + one overall level correction", "0.56", "0.69", "×1.42", "81%", "8/10"],
       ["pre-1200 gradient boosting, strict ex-ante", "0.34", "0.69", "×1.57", "66%", "0/10"]],
      "The no-leakage forecast: every coefficient estimated on 800–1200 transitions only, then "
      "iterated 1200→1500. The ‘level correction’ is a single scalar shift applied to all "
      "cities — it cannot reorder the hierarchy.")
para("The strict forecast under-predicts the overall level (mean error −0.23 log): the centuries "
     "after 1200 were collectively faster than those before, and a law trained before 1200 "
     "cannot know that. But a common level miss shifts every city equally — it costs level-R² "
     "while leaving the predicted hierarchy untouched, which is why the rank correlation (0.69) "
     "and the top-10 identification (8/10) exactly match the within-window fit, and why one "
     "overall scalar restores R² = 0.56. The 1500 hierarchy was predictable not merely from 1200 "
     "data, but from the growth regime of the pre-1200 world. (A gradient booster trained on the "
     "same pre-1200 transitions degrades badly out of its training range — trees cannot "
     "extrapolate levels — which is itself evidence that the linear persistence-plus-water "
     "structure, not model flexibility, carries the forecast.)")
h2("8.2 Is 0.56 high or low? The information ceiling")
para("An R² of 0.56 invites the objection: barely more than half the variation — surely a better "
     "model could do better? The objection assumes the missing 44% is signal waiting for a "
     "cleverer specification. That is testable. I gave flexible machine-learning models "
     "(random forest, gradient boosting) every scrap of 1200-vintage information in the data — "
     "size in 1200, size in 800, coordinates, elevation, river, coast, sea basin, every "
     "privilege already granted by 1200, and the coverage flags themselves — and measured "
     "out-of-sample accuracy two ways: ordinary five-fold CV over cities, and spatially blocked "
     "CV (whole 2°×2° map cells held out together), so that raw coordinates cannot let the "
     "trees simply recognise a region and recall how big its cities ended up. Table 5 is the "
     "result.")
table(["model, all restricted to information available in 1200", "random 5-fold R²", "spatially blocked R²"],
      [["the 3-parameter law (size + water + shocks)", "0.56", "0.56"],
       ["linear (ridge), all features", "0.59", "0.58"],
       ["random forest, all features", "0.61", "0.58"],
       ["gradient boosting, all features", "0.62", "0.60"]],
      "The information ceiling. No model, however flexible, extracts much more from 1200-vintage "
      "information than the 3-parameter law already does — and the flexible models' small edge "
      "shrinks further once cross-validation is spatially blocked.")
para("The kitchen sink buys about four to six points of R², and the model's own accounting says "
     "where they come from: 78% of the gradient booster's predictive weight sits on 1200 size "
     "alone, most of the rest on raw coordinates and elevation (which proxy the coming "
     "Atlantic–North Sea shift, the pattern of §7), and under 1% on the privileges. Roughly "
     "0.60 is the ceiling of what could have been known in 1200; the law reaches 0.56 of it "
     "with three parameters. The residual ~40% was not knowable — it is three centuries of "
     "plague draws, fires, sieges, trade-route accidents, and banking booms, plus the rounding "
     "and imputation error documented in §2.")
para("Two calibrations make the 0.56 easier to judge. First, the horizon ladder: predicting 1500 "
     "sizes from 1400 sizes — one century ahead, with perfect knowledge of the starting point — "
     "achieves only R² = 0.76; from 1300, 0.65; from 1200, 0.56. Unpredictability compounds "
     "century by century, and the 300-year figure is exactly on that decay curve. There is no "
     "specification that beats the horizon; even God's own 1400 census caps out at 0.76 for "
     "1500. Second, scale: 300 years is roughly twelve human generations. Social science "
     "celebrates R² ≈ 0.1–0.25 for transmitting economic status across a single generation. "
     "Predicting individual cities across twelve, from three numbers, at 0.56, is not a weak "
     "result — it is an extraordinary amount of determinism, which is precisely the paper's "
     "claim of path dependence.")
para("The deeper point is that 0.56 is not a grade on the model; it is a measurement of the "
     "system. The paper's central decomposition — roughly two-thirds of the hierarchy locked in, "
     "roughly a third unexplained — and the prediction R² are the same fact in different units. If some "
     "specification pushed R² to 0.9, it would not vindicate the analysis; it would falsify "
     "Finding 3, because a system whose growth is 80% noise (κ/σ = 0.21) cannot be 90% "
     "predictable three centuries out. The equation predicts as well as the world it describes "
     "permits, and the shortfall is itself the measured quantity: the share of urban fate that "
     "was not knowable in advance — the share this paper labels unexplained and weakly "
     "persistent.")
h2("8.3 What the equation cannot do — and why that failure is informative")
para("Simulating the law forward from the real 1200 sizes (400 runs) reproduces the persistence of "
     "the hierarchy and the general shape of its evolution. It fails in exactly one direction: the "
     "real top tier concentrated more than randomness predicts. The Zipf exponent — a standard "
     "measure of how dominant the largest cities are, lower = more concentrated — fell from 1.26 "
     "to 1.08 in the data, while the simulated law drifts toward mild equalization (1.22 by 1500). "
     "Figure 11 shows the divergence. That gap is the fingerprint of a force absent from the "
     "equation: agglomeration at the very top. Past a certain size, size itself attracts size — "
     "consistent with Dittmar's finding that Zipf's law for European cities emerges as the upper "
     "tail thickens toward the modern era.")
fig("fig_zipf_evolution.png",
    "Concentration of the urban hierarchy (Zipf exponent; lower = largest cities more dominant). "
    "Real cities (red) concentrate faster than the fitted random-growth law predicts (grey band).")
para("Medieval city size followed a near-random walk anchored on geography and swept by common "
     "shocks — a stochastic law, not a factor recipe — plus a modest agglomeration pull at the "
     "top that is the one thing the random walk cannot fake.", "Finding 3:")

# ================================================================ 9
h1("9. Do the Results Survive the Data's Own Weaknesses?")
h2("9.1 The Bairoch swap")
para("Every number above comes from the Buringh panel. Since medieval populations are "
     "reconstructions, I re-ran the core analyses with Bairoch's (1988) estimates swapped in — "
     "same cities, same geography, same privilege dates and coverage universes; only the "
     "population numbers change. The two compilations overlap in their sources (§2), so this is "
     "a test of sensitivity to the reconstruction choices, not a fully independent replication — "
     "but it is exactly the comparison that matters for the results that could be artifacts of "
     "one scholar's imputation scheme. Figure 12 shows the raw agreement (r = 0.96 on 1500 log "
     "sizes, 511 common cities); Table 6 and Figure 13 place the headline quantities side by "
     "side.")
fig("fig_source_agreement.png",
    "Agreement between the two population reconstructions on 1500 city sizes "
    "(r = 0.96, n = 511).")
table(["result", "Buringh (2021)", "Bairoch (1988)"],
      [["growth predictable from size? (R², 1300→1500)", "0.05", "0.30"],
       ["persistence of the hierarchy (r², 1300→1500)", "0.65", "0.54"],
       ["momentum share of 1500 size variance", "96%", "98%"],
       ["water premium per century (1300→1500 window)", "+9% (p = .008)", "+15% (p = .15)"],
       ["staple right: raw size gap (in-footprint)", "+127%", "+97%"],
       ["staple right: matched DiD", "−6% (53 treated)", "−41% (7 treated)"]],
      "Core results on the two population reconstructions, both under the corrected coverage "
      "universes.")
fig("fig_robustness_bars.png",
    "The same qualitative findings on both reconstructions: persistence dominates, water is "
    "positive, the raw staple gap collapses under the causal test.")
para("Every qualitative conclusion replicates: century growth is largely unpredictable, inherited "
     "size dominates, water carries a positive premium, and the raw staple advantage collapses "
     "(indeed reverses) under the timing-aware comparison. The honest differences run in expected "
     "directions: Bairoch is thin before 1300 and skews to larger, better-documented towns, so it "
     "shows more mean-reversion and its staple test rests on only seven treated towns. Nothing "
     "rests on which scholar's numbers one prefers.")
h2("9.2 Imputation and the small-town floor")
para("Buringh's own metadata says the series combines earlier demographic compilations with new "
     "estimates and city/time-specific imputations, and §2 quantified it: 81% of the city-year "
     "values in the 1200–1500 analysis sample are flagged imputed or proxied (95% in 1200, "
     "falling to 69% by 1500). This cannot be waved away, and it cannot be fully purged either: "
     "restricting to purely non-imputed observations leaves about 40 usable cities — far too few "
     "to re-estimate anything. Three mitigations are available. First, the medieval urban "
     "hierarchy is not controversial at the top, and the results are rank-driven. Second, the "
     "Bairoch swap above changes the imputation scheme wholesale and changes no conclusion. "
     "Third, rounding-to-thousands noise is worst for the smallest towns, so Table 7 raises the "
     "population floor from 1,000 to 5,000 and 10,000: persistence and the water premium's "
     "point estimate barely move (significance fades with sample size, as it must), and the "
     "staple DiD stays at or below zero throughout.")
table(["population floor", "n (1200 & 1500)", "persistence r²", "Gibrat R² (growth~size)",
       "water premium", "staple DiD"],
      [["1,000", "1,092", "0.56", "0.03", "+0.14 (p < .001)", "−0.08 (65 treated)"],
       ["5,000", "153", "0.52", "0.00", "+0.13 (p = .18)", "−0.06 (17 treated)"],
       ["10,000", "59", "0.49", "0.01", "+0.14 (p = .36)", "−0.37 (3 treated)"]],
      "Threshold robustness. Rounded populations make small-town growth noisy; the core "
      "quantities are stable as the floor rises and samples shrink.")

# ================================================================ 10
h1("10. What Is Actually New Here")
para("Each ingredient of this paper has a literature behind it; the contributions are specific:")
bullet("Privileges get a timing-aware causal test with explicit coverage discipline — and the "
       "null generalizes across Europe. What is new is difference-in-differences and stacked "
       "event-study estimates around the grant dates of six institutional treatments from four "
       "sources (Viabundus staples and fairs; Städtebuch charters and market rights; Bosker "
       "communes Europe-wide; Wahl participative institutions), each restricted to the universe "
       "where absence is observable. The event studies show privileged towns rising for two "
       "centuries before the grant with no break after it — in Germany, Italy, and France "
       "alike. For the well-powered tests the nulls are informative: staple effects above +10%, "
       "fair above +14%, commune above +7%, and participative-institution effects above +8% are "
       "excluded, against naive gaps of +40%, +26%, +21%, +5%; for charters and market rights "
       "the nulls are honest but imprecise (MDEs of +41%/+51%). Prior work (e.g., Bosker et "
       "al. 2013) showed institution coefficients shrink under fixed effects; this is a direct "
       "before/after test of the grants themselves.", "1.")
bullet("The urban hierarchy's origins are decomposed, with the decomposition's own limits "
       "stated. The full-sample Shapley split of predictive R² — 48% inherited (1200) size, 19% "
       "deep (800 AD) origin, 2.6% geography conditional on those sizes, 31% unexplained — is "
       "descriptive accounting; geography's small share is a statement about what remained for "
       "geography to do after 1200, not about its total historical role. Institutions, measured "
       "on their observable universe, carry ~12 points of variance as markers and ~4 points "
       "incremental to geography-plus-size. The residual-persistence test (autocorrelation "
       "−0.10/−0.16) rules out a stable omitted factor; the residual itself is unexplained and "
       "weakly persistent — transient shocks plus measurement noise, bounded above by ‘luck.’ "
       "Persistence quantified to 800 AD (r² = 0.40 across 700 years) puts a number on ‘deep "
       "roots’ claims usually made qualitatively.", "2.")
bullet("An explicit implied-size benchmark turns ‘over-performance’ into a measured quantity per "
       "city (Hamburg 9.7×, Mainz 0.5×), and the residuals line up on one axis: water.", "3.")
bullet("The summary equation is validated as a predictor with the temporal leak closed: "
       "coefficients estimated solely on 800→1200 transitions, locked, still forecast the 1500 "
       "hierarchy at R² = 0.50 strict (0.56 with one scalar level correction), rank correlation "
       "0.69, 8/10 of the largest cities — against R² = 0.02 for growth — and sit near the "
       "information ceiling (flexible models with every 1200-vintage feature: ≈0.60 under "
       "spatially blocked CV). The one systematic failure (real concentration outrunning the "
       "simulated law, Zipf 1.08 vs 1.22) isolates top-end agglomeration as the single force a "
       "random walk cannot mimic.", "4.")
bullet("The Black Death result extends Jedwab, Johnson & Koyama's recovery-toward-fixed-factors "
       "finding — and the mechanism is confronted with actual arrival data. Within this sample "
       "the operative fixed factor was navigable water: the coastal premium is negative in the "
       "plague century and strongly positive in the recovery century, with post-plague winners "
       "disproportionately on water. Matching Biraben's digitized first-wave outbreak records "
       "(1347–52) confirms the exposure geography — coastal and river cities were 6–7 points "
       "more likely to record an outbreak, size-controlled — while showing candidly that the "
       "damage-and-rebound channel cannot be identified from chronicle records at century "
       "resolution.", "5.")

# ================================================================ 11
h1("11. What This Means")
para("The story in which a wise ruler's grant of staple or market rights created a great city "
     "reverses the arrow the data show. Rulers granted privileges to places that were already "
     "rising — the event studies put the ascent's start at least two centuries before the "
     "typical grant — usually because water access or Roman-era roots had put them on the map "
     "long before. The privilege was recognition of success in progress, not its cause; no "
     "privilege shows a growth break at its grant, and for the best-measured privileges any "
     "effect beyond +10–14% is ruled out. None of this makes charters unimportant to medieval "
     "life — they shaped law, liberty, and identity — but they do not explain who grew.")
para("What does explain who grew is boring, which is the point. A city's fortune was set by where "
     "it sat in the landscape of trade — above all, whether goods could reach it by water — and "
     "by how early it had established itself, because size carried itself forward across "
     "centuries at a rate of 0.915 per century. Around that skeleton, individual fates were "
     "shuffled by shocks no equation predicts: a residual five times larger than the systematic "
     "pull. Policy operated at the margins of a system whose destiny was mostly mapped before "
     "the policies existed. If there is a lesson for the long run of cities, it is humility about "
     "recipes: the levers everyone argued about moved almost nothing, the ones nobody could move "
     "— rivers, coasts, head starts — moved almost everything, and a large remainder was "
     "shocks no one could foresee and no ledger recorded.")

# ================================================================ refs
h1("References")
for ref in [
    "Bairoch, P., Batou, J., & Chèvre, P. (1988). La population des villes européennes de 800 à "
    "1850. Geneva: Droz.",
    "Bosker, M., Buringh, E., & van Zanden, J. L. (2013). From Baghdad to London: Unraveling urban "
    "development in Europe, the Middle East, and North Africa, 800–1800. Review of Economics and "
    "Statistics, 95(4), 1418–1437.",
    "Buringh, E. (2021). The population of European cities from 700 to 2000. Research Data Journal "
    "for the Humanities and Social Sciences, 6(1), 1–18.",
    "Cantoni, D., Mohr, C., & Weigand, M. (2020). Princes and townspeople: Dated town charters and "
    "market rights from the Deutsches Städtebuch.",
    "Dittmar, J. (2011). Cities, markets, and growth: The emergence of Zipf's law. Working paper.",
    "Gabaix, X. (1999). Zipf's law for cities: An explanation. Quarterly Journal of Economics, "
    "114(3), 739–767.",
    "Gibrat, R. (1931). Les inégalités économiques. Paris: Sirey.",
    "Holterman, B., et al. Viabundus: Map of premodern European transport and mobility, 1350–1650.",
    "Israeli, O. (2007). A Shapley-based decomposition of the R-square of a linear regression. "
    "Journal of Economic Inequality, 5(2), 199–212.",
    "Jedwab, R., Johnson, N. D., & Koyama, M. (2024). Medieval cities through the lens of urban "
    "economics.",
    "Krauer, F., & Schmid, B. V. (2022). Mapping the plague through natural language processing "
    "(digitized Biraben/Sticker outbreak data). Zenodo, doi:10.5281/zenodo.6587267.",
    "Wahl, F. (2015). Participative political institutions in pre-modern Europe: Introducing a "
    "new database. Historical Methods, 48(3), 110–124.",
    "Shorrocks, A. F. (1982). Inequality decomposition by factor components. Econometrica, 50(1), "
    "193–211.",
]:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(9.5)

doc.save(DOC)
print(f"saved {DOC}")
